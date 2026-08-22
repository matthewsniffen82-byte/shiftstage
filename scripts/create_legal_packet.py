from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


VIOLET = "6D28D9"
DEEP_VIOLET = "3B0764"
CYAN = "0891B2"
INK = "17131F"
MUTED = "5F5A67"
LIGHT_VIOLET = "F3EEFF"
LIGHT_CYAN = "ECFEFF"
LIGHT_GRAY = "F5F4F7"
WHITE = "FFFFFF"
BLACK = "07070A"
GREEN = "166534"
AMBER = "92400E"
RED = "991B1B"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge_data in edges.items():
        tag = f"w:{edge_name}"
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_hyperlink(paragraph, text: str, url: str, color: str = CYAN):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_repeat_on_every_page(section, value: bool) -> None:
    section.header.is_linked_to_previous = value
    section.footer.is_linked_to_previous = value


def add_header_footer(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("MYDANCR  /  ATTORNEY REVIEW DRAFT")
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(VIOLET)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "D7C8FF")
    borders.append(bottom)
    p_pr.append(borders)

    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(6.7))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(5.55)
    table.columns[1].width = Inches(1.15)
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = left.add_run("Confidential working draft · Not approved for publication")
    run.font.name = "Aptos"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    right = table.cell(0, 1).paragraphs[0]
    set_page_number(right)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 0, 0, 0, 0)
        set_cell_border(cell, top={"val": "single", "sz": "6", "color": "DDD9E4"})


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    add_header_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Title", 30, INK, 0, 10),
        ("Subtitle", 13, MUTED, 0, 8),
        ("Heading 1", 19, DEEP_VIOLET, 14, 7),
        ("Heading 2", 13, VIOLET, 10, 5),
        ("Heading 3", 10.5, CYAN, 7, 3),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1", "Heading 2"} else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(9.2)
        style.paragraph_format.left_indent = Inches(0.22)
        style.paragraph_format.first_line_indent = Inches(-0.15)
        style.paragraph_format.space_after = Pt(3)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("MYDANCR")
    r.font.name = "Aptos Display"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(VIOLET)

    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = band.cell(0, 0)
    set_cell_shading(cell, BLACK)
    set_cell_margins(cell, top=420, start=360, bottom=420, end=360)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "18", "color": VIOLET},
        bottom={"val": "single", "sz": "18", "color": CYAN},
        start={"val": "single", "sz": "10", "color": DEEP_VIOLET},
        end={"val": "single", "sz": "10", "color": DEEP_VIOLET},
    )
    title = cell.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("LEGAL REVIEW\nPACKET")
    run.font.name = "Aptos Display"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(WHITE)
    subtitle = cell.add_paragraph()
    sr = subtitle.add_run("Factual Product Map + Draft Agreements and Policies")
    sr.font.name = "Aptos"
    sr.font.size = Pt(13)
    sr.font.bold = True
    sr.font.color.rgb = RGBColor.from_string("C4F5FF")

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    values = [
        ("Prepared for", "Specialist adult-entertainment counsel"),
        ("Product", "MyDancr · mobile-first web application"),
        ("Review date", "August 21, 2026"),
        ("Status", "Attorney review draft · not approved for publication"),
    ]
    for row, (label, value) in zip(meta.rows, values):
        row.cells[0].width = Inches(1.45)
        row.cells[1].width = Inches(5.15)
        for cell in row.cells:
            set_cell_margins(cell, 100, 120, 100, 120)
            set_cell_border(cell, bottom={"val": "single", "sz": "5", "color": "DDD9E4"})
        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(label.upper())
        lr.font.bold = True
        lr.font.size = Pt(7.5)
        lr.font.color.rgb = RGBColor.from_string(VIOLET)
        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(9.5)
        vr.font.color.rgb = RGBColor.from_string(INK)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(18)
    note.paragraph_format.space_after = Pt(0)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "This packet is a product-informed drafting aid for licensed counsel. It is not legal advice, a legal opinion, or a substitute for counsel’s jurisdiction-specific review."
    )
    nr.font.name = "Aptos"
    nr.font.size = Pt(9)
    nr.font.italic = True
    nr.font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_page_break()


def parse_inline(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|https?://\S+)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Cascadia Mono"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string(DEEP_VIOLET)
        else:
            clean = token.rstrip(".,);]")
            suffix = token[len(clean):]
            add_hyperlink(paragraph, clean, clean)
            if suffix:
                paragraph.add_run(suffix)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_callout(doc: Document, label: str, text: str, kind: str) -> None:
    palette = {
        "COUNSEL": (LIGHT_VIOLET, VIOLET),
        "PRODUCT FACT": (LIGHT_CYAN, CYAN),
        "IMPLEMENTATION": (LIGHT_GRAY, MUTED),
        "HIGH RISK": ("FEF2F2", RED),
        "DECISION": ("FFF7ED", AMBER),
    }
    fill, accent = palette.get(kind, (LIGHT_GRAY, MUTED))
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 120, 160, 120, 160)
    set_cell_border(
        cell,
        start={"val": "single", "sz": "22", "color": accent},
        top={"val": "single", "sz": "4", "color": accent},
        bottom={"val": "single", "sz": "4", "color": accent},
        end={"val": "single", "sz": "4", "color": accent},
    )
    p = cell.paragraphs[0]
    r = p.add_run(f"{label.upper()}  ")
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(accent)
    parse_inline(p, text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.75)
    table.columns[1].width = Inches(4.9)
    hdr = table.rows[0]
    hdr.cells[0].text = "ITEM"
    hdr.cells[1].text = "CURRENT DRAFT POSITION / REVIEW POINT"
    set_repeat_table_header(hdr)
    for cell in hdr.cells:
        set_cell_shading(cell, DEEP_VIOLET)
        set_cell_margins(cell, 90, 110, 90, 110)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor.from_string(WHITE)
    for index, (key, value) in enumerate(rows):
        row = table.add_row()
        keep_row_together(row)
        if index % 2:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for cell in row.cells:
            set_cell_margins(cell, 90, 110, 90, 110)
            set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": "DDD9E4"})
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        kp = row.cells[0].paragraphs[0]
        kr = kp.add_run(key)
        kr.bold = True
        kr.font.size = Pt(8.5)
        vp = row.cells[1].paragraphs[0]
        parse_inline(vp, value)
        for run in vp.runs:
            run.font.size = Pt(8.5)


def add_flow_table(doc: Document, rows: list[tuple[str, str, str, str, str]]) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.23)
    section.footer_distance = Inches(0.23)
    add_header_footer(section)
    doc.add_heading("End-to-end data and money flow matrix", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["TRIGGER", "DATA / VALUE", "PROCESSING", "RECIPIENT / OUTPUT", "CONTROL / RECORD"]
    widths = [1.5, 2.1, 2.45, 2.1, 2.15]
    for i, (header, width) in enumerate(zip(headers, widths)):
        table.columns[i].width = Inches(width)
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, DEEP_VIOLET)
        set_cell_margins(cell, 70, 80, 70, 80)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(6.8)
            run.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    for index, values in enumerate(rows):
        row = table.add_row()
        keep_row_together(row)
        if index % 2:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for i, value in enumerate(values):
            cell = row.cells[i]
            set_cell_margins(cell, 70, 80, 70, 80)
            set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": "D7D2DC"})
            p = cell.paragraphs[0]
            parse_inline(p, value)
            for run in p.runs:
                run.font.size = Pt(7.3)
    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    portrait.orientation = WD_ORIENT.PORTRAIT
    portrait.page_width, portrait.page_height = portrait.page_height, portrait.page_width
    portrait.top_margin = Inches(0.65)
    portrait.bottom_margin = Inches(0.65)
    portrait.left_margin = Inches(0.72)
    portrait.right_margin = Inches(0.72)
    portrait.header_distance = Inches(0.28)
    portrait.footer_distance = Inches(0.28)
    add_header_footer(portrait)


def render_markdown(doc: Document, source: str) -> None:
    lines = source.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped:
            i += 1
            continue
        if stripped == "---PAGE---":
            doc.add_page_break()
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("[CALLOUT:"):
            match = re.match(r"\[CALLOUT:([^|\]]+)\|([^\]]+)\]\s*(.*)", stripped)
            if not match:
                raise ValueError(f"Invalid callout syntax: {stripped}")
            add_callout(doc, match.group(2), match.group(3), match.group(1))
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(stripped[2:])
            r.italic = True
            r.font.color.rgb = RGBColor.from_string(MUTED)
            p_pr = p._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "14")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), CYAN)
            borders.append(left)
            p_pr.append(borders)
        else:
            paragraph_lines = [stripped]
            j = i + 1
            while j < len(lines):
                candidate = lines[j].strip()
                if not candidate or candidate.startswith(("#", "- ", "> ", "[CALLOUT:", "---PAGE---")) or re.match(r"^\d+\.\s", candidate):
                    break
                paragraph_lines.append(candidate)
                j += 1
            p = doc.add_paragraph()
            parse_inline(p, " ".join(paragraph_lines))
            i = j - 1
        i += 1


def add_source_list(doc: Document, sources: list[tuple[str, str]]) -> None:
    doc.add_heading("Selected primary authorities and provider materials", level=1)
    intro = doc.add_paragraph(
        "Counsel should confirm the applicability, current version, and jurisdictional scope of every authority before publication. This list supports issue spotting; it is not an opinion that any statute necessarily applies."
    )
    intro.paragraph_format.space_after = Pt(8)
    for title, url in sources:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{title}: ").bold = True
        add_hyperlink(p, url, url)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: create_legal_packet.py SOURCE.md OUTPUT.docx")
    source_path = Path(sys.argv[1]).resolve()
    output_path = Path(sys.argv[2]).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    set_document_defaults(doc)
    add_cover(doc)

    source = source_path.read_text(encoding="utf-8")
    before_flow, after_flow = source.split("\n[[FLOW_MATRIX]]\n", 1)
    render_markdown(doc, before_flow)

    flow_rows = [
        ("Customer signs up", "Email, city, password/auth metadata", "Supabase Auth creates user; app assigns customer role and private profile", "Customer dashboard", "Confirmation, rate limit, row-level access"),
        ("Dancer signs up", "Email, city, eligibility/profile fields", "Auth + draft dancer profile; contractual eligibility; staged media and club-affiliation onboarding", "Private dancer workspace", "Email confirmation; acceptance version; account state; audit"),
        ("Dancer uploads media", "Image/video, metadata, moderation signals", "Signature/type/size checks; transform/watermark; automated and/or human moderation", "Approved public media or private review state", "Moderation record; storage path; rejection reason"),
        ("Dancer submits", "Profile, accepted eligibility, approved media", "Account/media checks; first venue affiliation remains required", "Pending/approved state", "Approval review; notification"),
        ("Dressing-room NFC tap", "Tag token, dancer account, venue, device/IP audit", "Secure tag validation; affiliation/working session creation", "Venue affiliation + time-limited Working Now", "Tap event; expiry; cooldown; audit"),
        ("Dancer posts shift", "Approved affiliated club, date/time", "Server restricts venue to active approved affiliation", "Public upcoming schedule", "Shift record; notification; analytics"),
        ("Customer discovers", "City/filter, coarse session/device data", "Public eligibility and ranking queries", "Dancer/club/TV results", "View and ranking events"),
        ("Customer follows/saves", "Account or anonymous session token", "Relationship/preference record", "Personalized dashboard/notifications", "Follow, favorite, going, saved-deal records"),
        ("Club publishes deal", "Admission/line-access terms + accepted offer-specific Deal Order", "Authorization, deal-policy, fee, dates, and active-deal validation", "Public Club Deal", "Deal Order + fee snapshot + deal + club/admin activity"),
        ("Customer selects deal", "Deal, source page, attributed dancer/shift when eligible", "Signed redemption token is created and attribution locked", "Cashier redemption screen", "Redemption + attribution audit"),
        ("Cashier NFC tap", "Cashier tag, token, session/device/IP", "Venue/deal/status/expiry/duplicate checks; server confirmation", "Verified redemption", "NFC tap, redemption event, fraud indicators"),
        ("Commission accrues", "Profile-originated verified redemption + that deal’s negotiated referral fee", "Monthly event 1–9: 30%; 10–24: 40%; 25+: 50%", "Independent dancer commission eligible for NATS export", "Immutable source + Deal Order/fee + monthly ordinal + 3,000/4,000/5,000 bp + policy version"),
        ("Commission exported", "Verified NATS affiliate ID + exact eligible USD amount", "Durable NATS outbox exports once; ambiguous responses require reconciliation", "Too Much Media/NATS affiliate ledger", "Linkage + attempt/status + provider reference + reversal; no external account credentials in app DB"),
        ("Club settles weekly", "Weekly itemized statement + accepted ACH authorization/instructions", "ACH debit or credit, return/retry, and receivable reconciliation", "MyDancr club receivable settled", "Statement/item + ACH trace/status/return; separate from dancer commission"),
        ("Notification event", "Recipient, event type, preference, channel", "In-app plus optional push/email delivery", "Account holder", "Notification/delivery status; opt-out preference"),
        ("Report/support request", "Target, category, description, attachments/contact", "Triage, moderation, escalation, response", "Admin/support queue and reporter updates", "Case/thread/messages/actions"),
        ("DMCA notice", "Claimant, work, location, statements, signature", "Completeness review; removal/counter-notice/restore workflow", "Claimant, uploader, administrator", "DMCA case, counter, strike, actions"),
        ("Account deletion", "Authenticated request", "Auth user deletion and cascade/cleanup subject to legal holds", "Access disabled; public content removed", "Minimum audit, finance, fraud, DMCA, tax/legal records retained as approved"),
    ]
    add_flow_table(doc, flow_rows)
    render_markdown(doc, after_flow)

    doc.core_properties.title = "MyDancr Legal Review Packet"
    doc.core_properties.subject = "Factual product map and draft agreements/policies for attorney review"
    doc.core_properties.author = "MyDancr product team"
    doc.core_properties.keywords = "MyDancr, legal review, Arizona pilot, terms, privacy, dancer, club, NFC, Club Deals, ACH, NATS, commissions"
    doc.core_properties.comments = "Attorney review draft; not approved for publication"
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
