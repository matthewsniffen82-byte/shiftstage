from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "171717"
MUTED = "555555"
LIGHT_GRAY = "F2F4F7"


def normalize_ascii(text: str) -> str:
    replacements = {
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u2212": "-",
        "\u00b7": "|",
        "\u00a0": " ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def extract_part(source: str, roman: str) -> str:
    marker = f"# Part {roman} "
    start = source.index(marker)
    end = source.find("\n---PAGE---", start)
    if end == -1:
        end = len(source)
    return source[start:end].strip()


def replace_section(text: str, heading: str, next_heading: str | None, body: str) -> str:
    start = text.index(heading)
    if next_heading is None:
        end = len(text)
    else:
        end = text.index(next_heading, start)
    return text[:start] + body.strip() + "\n\n" + text[end:]


def clean_part(text: str, new_heading: str) -> str:
    lines = normalize_ascii(text).splitlines()
    lines[0] = new_heading
    lines = [line for line in lines if not line.startswith("[CALLOUT:")]
    return "\n".join(lines).strip()


def build_source(source: str) -> str:
    terms = clean_part(extract_part(source, "II"), "# Document 1 - Terms of Service")
    privacy = clean_part(extract_part(source, "III"), "# Document 2 - Privacy Policy")
    dancer = clean_part(extract_part(source, "IV"), "# Document 3 - Dancer Agreement")
    club = clean_part(extract_part(source, "V"), "# Document 4 - Club Agreement")
    media = clean_part(extract_part(source, "VI"), "# Schedule A to the Dancer Agreement - Content and Media Consent/License")
    acceptable = clean_part(extract_part(source, "VII"), "# Schedule B to the Dancer Agreement - Acceptable Use and Prohibited Conduct Rules")
    eligibility = clean_part(extract_part(source, "VIII"), "# Schedule C to the Dancer Agreement - Account Eligibility and Venue Verification")
    dmca = clean_part(extract_part(source, "IX"), "# Document 5 - DMCA and Takedown Policy")
    deal = clean_part(extract_part(source, "X"), "# Schedule A to the Club Agreement - Club Deal and NFC Redemption Terms")
    commission = clean_part(extract_part(source, "XI"), "# Schedule D to the Dancer Agreement - Dancer Commission and Too Much Media/NATS Settlement Terms")

    terms = replace_section(
        terms,
        "## 6. Club Deals, NFC, and attribution",
        "## 7. Club payments and dancer commission settlement",
        """
## 6. Club Deals, NFC, and attribution

Club Deals are conditional club offers, not cash, stored value, reservations, tickets, or guaranteed admission. The club controls admission and must honor a valid deal subject to disclosed limits, capacity, lawful age and identification requirements, dress code, safety, hours, and law. A deal is verified only through the authorized server and cashier NFC flow. Screenshots, copied tokens, expired links, or unconfirmed selections are invalid.

Attribution may record the club page, dancer profile, shift or deal surface, browser session, deal selection, and NFC confirmation that led to a verified redemption. MyDancr may use this attribution to create weekly club statements, reconcile ACH, and calculate dancer commissions. Users may not manipulate source attribution, create duplicate redemptions, simulate taps, or interfere with tags.

MyDancr does not use geofences or continuous or background location tracking to determine club attendance, dancer presence, or deal redemption. An in-club verification or redemption is established only by an intentional, server-validated tap of the applicable official MyDancr NFC tag.
""",
    )
    terms = replace_section(
        terms,
        "## 11. Reports, enforcement, and appeals",
        "## 12. Intellectual property and DMCA",
        """
## 11. Reports, enforcement, and appeals

MyDancr may investigate reports, preserve evidence, restrict content, disable features, suspend accounts, terminate access, revoke tags, void fraudulent redemptions, and notify affected parties or authorities as permitted or required. When appropriate, MyDancr will provide notice of the action, the principal reason, and a reasonable method to request review. Immediate action may occur for safety, minors, trafficking, credible threats, fraud, sanctions, legal process, or system integrity.
""",
    )
    terms = replace_section(
        terms,
        "## 15. Limitation of liability",
        "## 16. Indemnity",
        """
## 15. Limitation of liability

To the maximum extent permitted by law, Company and its affiliates, officers, employees, contractors, and service providers will not be liable for indirect, incidental, special, consequential, exemplary, or punitive damages, or for lost profits, revenue, goodwill, data, business opportunity, or expected attendance, even if advised that such damages were possible. This exclusion does not apply where prohibited by law.

To the maximum extent permitted by law, Company's aggregate liability arising from the Service or these Terms will not exceed the greater of one hundred U.S. dollars or the amount the claimant paid directly to Company for the Service during the twelve months before the event giving rise to the claim. A separate signed Club Agreement may establish a different commercial cap. Nothing in these Terms limits liability that cannot lawfully be limited, including liability to the extent caused by fraud, willful misconduct, or gross negligence where applicable law so requires.
""",
    )
    terms = replace_section(
        terms,
        "## 16. Indemnity",
        "## 17. Termination",
        """
## 16. Indemnity

To the extent permitted by law, a dancer, club, or other business User will defend, indemnify, and hold harmless Company and its affiliates, officers, employees, and contractors from third-party claims, damages, judgments, penalties, costs, and reasonable attorneys' fees arising from that User's content, club or account authority, offers, unlawful conduct, licensing or tax failures, infringement, or material breach of these Terms or an incorporated agreement. Company will provide reasonable notice and cooperation and will not settle a covered claim in a manner that admits the indemnifying party's fault or imposes a nonmonetary obligation on it without consent. This clause does not require a User to indemnify Company for Company's own fraud, willful misconduct, or gross negligence.
""",
    )
    terms = replace_section(
        terms,
        "## 18. Governing law and disputes",
        "## 19. General",
        """
## 18. Governing law and disputes

These Terms are governed by Arizona law, without regard to conflict-of-law rules, except to the extent federal law or a User's nonwaivable local consumer law applies. Before filing a claim, the parties will make a good-faith effort for thirty days to resolve it through written notice describing the claim and requested relief. Unless applicable law permits a different forum, claims will be brought in the state courts located in Maricopa County, Arizona, or the United States District Court for the District of Arizona. Either party may seek eligible small-claims relief or immediate equitable relief for security, intellectual-property, confidentiality, or system-integrity harm. A separate Club Agreement may provide different negotiated dispute terms.
""",
    )
    terms = replace_section(
        terms,
        "## 19. General",
        None,
        """
## 19. General

These Terms, the Privacy Policy, and each role- or feature-specific agreement accepted by the User are the entire agreement regarding the Service. Specific terms control over general terms for the relevant feature. Notices may be delivered electronically to the account email or through the Service. Users consent to electronic records and communications. Company may assign these Terms to an affiliate or successor in connection with a reorganization, financing, merger, acquisition, or sale of assets. A User may not assign an account or these Terms without Company's written consent. Delay in enforcement is not a waiver. If a provision is unenforceable, it will be modified only to the minimum extent necessary and the remainder will continue. Neither party is liable for delay caused by events beyond reasonable control, except accrued payment obligations. Headings are for convenience only. The published legal page will identify Company's legal name, notice address, contact email, effective date, and current version.
""",
    )

    privacy = privacy.replace("### Location, NFC, device, and security data", "### Venue, NFC, device, and security data")
    privacy = privacy.replace(
        "For an eligible dancer check-in, latitude, longitude, accuracy, reading time, distance result, club, and verification expiry. For NFC and security events: tag/token digest, timestamp, club, browser session, device fingerprint or attributes, IP address, user agent, and audit payload. Public users do not receive raw dancer location.",
        "MyDancr stores club addresses and venue coordinates to display club information and link users to directions. For NFC and security events, MyDancr may process a tag or token digest, timestamp, club, browser session, limited device attributes, IP address, user agent, verification expiry, and audit payload. MyDancr does not create geofences, continuously or in the background track a User's location, or use device GPS to prove dancer presence, affiliation, or Club Deal redemption. Those in-club events are established only by an intentional, server-validated NFC tap. If a User chooses a directions or rideshare link, the independent map or rideshare provider may process the User's current location under its own terms and privacy notice; MyDancr may record only the destination link and ordinary click event needed to operate and measure the feature.",
    )
    privacy = privacy.replace(
        "Verify account eligibility representations, club authority, dancer affiliation, and time-limited presence.",
        "Verify account eligibility representations, club authority, dancer affiliation, and time-limited NFC-confirmed presence without geofencing or continuous location tracking.",
    )
    privacy = privacy.replace(
        "This policy does not govern independent club practices, third-party social sites, maps, rideshare providers, or a provider acting independently. The published version must identify the MyDancr controller/business legal entity, contact methods, actual production providers, and each provider's privacy role.",
        "This policy does not govern independent club practices, third-party social sites, maps, rideshare providers, or a provider acting independently. The published legal page identifies the MyDancr controller or business legal entity, contact methods, material production providers, and each provider's privacy role.",
    )
    privacy = privacy.replace(
        "The published policy must identify the contracted Too Much Media entity and link to the exact production terms/privacy notice before live data is sent. MyDancr must update the notice, contract/vendor record, and user acceptance before changing the commission-settlement provider or sending materially different data.",
        "Before live data is sent, the published policy will identify the contracted Too Much Media entity and link to the applicable production terms and privacy notice. MyDancr will update the notice, vendor record, and user acceptance when required before changing the commission-settlement provider or sending materially different data.",
    )
    privacy = privacy.replace(
        "Providers and users may be located in different jurisdictions. If MyDancr serves users outside the United States or sends data internationally, counsel must approve transfer mechanisms, controller/processor terms, localization rules, and rights notices before launch in that jurisdiction.",
        "Providers and Users may be located in different jurisdictions. When required, MyDancr will use approved transfer mechanisms, controller or processor terms, localization controls, and rights notices before serving a jurisdiction or transferring personal information internationally.",
    )
    privacy = replace_section(
        privacy,
        "## 9. Retention",
        "## 10. Security",
        """
## 9. Retention

MyDancr retains personal information only as long as reasonably necessary for the purposes described here and for applicable contracts, security, disputes, tax and accounting obligations, fraud prevention, DMCA administration, legal claims, and law. Account and public-profile information is ordinarily retained while the account is active and during a reasonable deletion and backup cycle. Finance, ACH, commission, tax, acceptance, security, affiliation, NFC, and redemption records may be retained through the applicable statutory, contractual, accounting, fraud, and limitations periods. Removed media may be retained in restricted form only when reasonably necessary for moderation, rights complaints, safety, legal hold, or defense of claims. Deletion requests do not require deletion of deidentified data or records that MyDancr must or is permitted to retain. Vendor copies are governed by contractual deletion and backup schedules.
""",
    )

    dancer = replace_section(
        dancer,
        "## 4. Working Now and presence data",
        "## 5. Profile and media obligations",
        """
## 4. Working Now and NFC presence events

Working Now requires an approved dressing-room NFC tap validated by the MyDancr server. The dancer authorizes collection of the limited tag, club, timestamp, browser session, device/IP, verification-expiry, and fraud-prevention data needed to validate the event and operate the time window. MyDancr does not use a geofence, continuously or in the background track the dancer's location, or use device GPS to establish Working Now. The dancer may not share a tag, simulate a tap, relay or copy a token, use automation, or ask another person to check in.

Working Now is time-limited and does not promise that the dancer remains continuously on premises or is available to any customer. The dancer should end an inaccurate status through available controls and notify support of a malfunction.
""",
    )
    dancer = replace_section(
        dancer,
        "## 10. Dancer acknowledgments",
        None,
        """
## 10. Dancer acknowledgments

The dancer acknowledges the customer-facing limits of badges and status; the absence of guaranteed earnings; the role of club admission and independent club rules; MyDancr's account-eligibility, affiliation, media-approval, NFC, and commission-calculation roles; Too Much Media/NATS's independent commission-settlement role; the absence of a separate age-verification or geofence-tracking service; the public nature of approved profile content; dynamic ranking signals; and the Terms, Privacy Policy, and Schedules A through D attached to this Agreement.
""",
    )
    dancer = dancer.replace(
        "Ordinary appeals follow the enforcement policy. Immediate restriction may precede notice where necessary.",
        "When appropriate, MyDancr will provide notice of the action, the principal reason, and a reasonable method to request review. Immediate restriction may precede notice where necessary for safety, fraud prevention, legal process, or system integrity.",
    )

    club = club.replace(
        "The club authorizes collection of tap, device/IP, tag, venue, timing, and fraud data necessary to operate affiliation, presence, and redemption. MyDancr may disable or rotate a tag for security, misuse, termination, or maintenance.",
        "The club authorizes collection of tap, device/IP, tag, venue, timing, and fraud data necessary to operate affiliation, time-limited presence, and redemption. MyDancr does not use geofences or continuous or background location tracking for those functions; the applicable intentional NFC tap is the verification event. MyDancr may disable or rotate a tag for security, misuse, termination, or maintenance.",
    )
    club = replace_section(
        club,
        "## 12. Commercial protections",
        None,
        """
## 12. Commercial protections

The club represents that it is duly organized or lawfully operated, has authority to enter this Agreement, maintains all licenses required for its premises and offers, and will comply with law. Each party will protect the other's nonpublic business, security, and personal information using reasonable safeguards and will use it only for this relationship.

The club will defend, indemnify, and hold harmless MyDancr and its affiliates, officers, employees, and contractors from third-party claims arising from the club's premises, admission, employment or contractor practices, licenses, offers, fulfillment, submitted content, taxes, or material breach, except to the extent caused by MyDancr's fraud, willful misconduct, or gross negligence. MyDancr will defend and indemnify the club from third-party claims that the unmodified MyDancr technology infringes a United States intellectual-property right, subject to notice, control of defense, and customary exclusions.

To the maximum extent permitted by law, neither party is liable for indirect, incidental, special, consequential, exemplary, or punitive damages, or lost profits, revenue, goodwill, data, or business opportunity. Except for accrued payment duties, confidentiality or security breach, indemnity obligations, infringement, fraud, willful misconduct, or gross negligence, each party's aggregate liability under this Agreement will not exceed the fees paid or payable to MyDancr under the affected Deal Orders during the twelve months before the event giving rise to the claim. The final signed Agreement or Deal Order may state a different negotiated cap.

This Agreement begins when electronically accepted by an authorized club representative and continues until terminated by either party on thirty days' written notice, unless a Deal Order states a committed term. Either party may terminate for an uncured material breach after ten days' notice, or immediately for illegality, safety risk, fraud, security compromise, insolvency, or loss of required license. Accrued payment, confidentiality, audit, dispute, indemnity, limitation, and record-retention obligations survive.

Arizona law governs, without regard to conflict-of-law rules. The parties will attempt in good faith for thirty days to resolve a dispute after written notice. Unless a signed Deal Order provides otherwise, exclusive venue is in the state courts in Maricopa County, Arizona, or the United States District Court for the District of Arizona. Notices may be electronic to the accepted account and legal contacts. This Agreement, its Schedules, and accepted Deal Orders are the entire agreement and may be amended only by an authorized electronic or written acceptance. A Deal Order controls for its specific commercial terms. Neither party may assign except to an affiliate or successor in a merger, reorganization, financing, or asset sale, and the club may not assign without ensuring the licensed premises and operator remain eligible.
""",
    )

    acceptable = acceptable.replace(
        "simulate a tap; manipulate location/time/device/session data; bypass rate limits;",
        "simulate a tap; manipulate tag, time, device, or session data; bypass rate limits;",
    )
    media = media.replace(
        "Promotional use outside the submitter's ordinary profile/feed context should require the consent scope approved by counsel and, where appropriate, a separate release.",
        "Promotional use outside the submitter's ordinary profile or feed context requires a separate consent or release when the accepted license does not clearly cover that use.",
    )
    media = media.replace(
        "Legal name is not public unless the dancer separately directs and counsel approves.",
        "Legal name is not public unless the dancer separately directs that disclosure and MyDancr lawfully approves it.",
    )
    media = media.replace(
        "The final policy should state whether MyDancr uses Content to train proprietary models; absent a clear opt-in and approved notice, this draft does not grant a training right.",
        "MyDancr will not use Content to train a proprietary model unless that use is separately disclosed and the submitter gives an express opt-in consent. This license does not grant a model-training right.",
    )
    media = media.replace(
        "At counsel's direction, this license does not impose proposed 18 U.S.C. §§ 2257/2257A producer-recordkeeping, labeling, custodian, or inspection duties and does not require a separate vendor age-verification process.",
        "This license does not impose a separate 18 U.S.C. §§ 2257/2257A producer-recordkeeping, labeling, custodian, or inspection program and does not require a separate vendor age-verification process.",
    )
    media = media.replace(
        "Counsel should re-review the allocation if MyDancr materially changes its accepted content, production role, or applicable jurisdiction.",
        "MyDancr will re-evaluate this allocation before materially changing its accepted content, production role, or applicable jurisdiction.",
    )
    eligibility = eligibility.replace(
        "NFC tags are scoped by type and club. Dressing-room tags support affiliation/presence; cashier tags support deal verification. Tokens are secret or hashed, tags have lifecycle states, and server confirmation is required. Tags should be installed in controlled areas and rotated on compromise.",
        "NFC tags are scoped by type and club. Dressing-room tags support affiliation and time-limited Working Now presence; cashier tags support deal verification. Tokens are secret or hashed, tags have lifecycle states, and server confirmation is required. Tags should be installed in controlled areas and rotated on compromise. MyDancr does not use geofences or continuous or background location tracking for these functions. The applicable intentional NFC tap is the sole in-club verification event.",
    )
    eligibility = eligibility.replace(
        "MyDancr should record only the evidence category, result, dates, reviewer, and decision unless counsel requires more.",
        "MyDancr will record only the evidence category, result, dates, reviewer, and decision unless additional information is legally necessary.",
    )
    deal = deal.replace(
        "Only a successful server response completes redemption.",
        "Only a successful server response completes redemption. MyDancr does not use geofences or continuous or background location tracking to infer redemption; the intentional cashier NFC tap is required.",
    )
    commission = commission.replace(
        "A payee must dispute a statement within the counsel-approved period and identify the specific event.",
        "A payee must dispute a statement within the period displayed in the accepted settlement schedule and identify the specific event.",
    )
    commission = commission.replace(
        "The final version must state MyDancr's export frequency and cutoff, Too Much Media/NATS settlement timing, minimum balance, fee allocation, rejected-payment handling, and final settlement.",
        "MyDancr will disclose the export frequency and cutoff, Too Much Media/NATS settlement timing, minimum balance, fee allocation, rejected-payment handling, and final-settlement process in the accepted settlement schedule or dashboard.",
    )
    club = club.replace(
        "The final order must state the weekly cutoff and timezone, statement delivery, ACH authorization or instructions, debit/due day, weekends/holidays, returns and retries, taxes, lawful late charges, failed-payment costs, reconciliation, record retention, and suspension rights.",
        "Each Deal Order states the weekly cutoff and timezone, statement delivery, ACH authorization or instructions, debit or due day, weekends and holidays, returns and retries, taxes, lawful late charges, failed-payment costs, reconciliation, record retention, and suspension rights.",
    )

    deal_order = """
# Schedule B to the Club Agreement - Club Deal Order and Weekly ACH Schedule

This Deal Order is incorporated into the Club Agreement between MyDancr and the club legal operator identified below. It applies only to the specific Club Deal described here. If this Deal Order conflicts with general Club Deal copy, this Deal Order controls the commercial terms.

## 1. Club and premises

**Legal operator:** [INSERT LEGAL ENTITY]

**DBA / public club name:** [INSERT DBA]

**Licensed premises address:** [INSERT ADDRESS]

**Authorized signer and title:** [INSERT NAME AND TITLE]

**MyDancr Deal Order number:** [INSERT ORDER NUMBER]

## 2. Club Deal

**Offer:** [INSERT EXACT ADMISSION OR LINE-ACCESS BENEFIT]

**Customer eligibility and limits:** [INSERT AGE/ID, HOURS, CAPACITY, DRESS CODE, FREQUENCY, EXPIRY, AND OTHER MATERIAL RESTRICTIONS]

**Offer start and end:** [INSERT DATES, TIMES, AND ARIZONA TIMEZONE]

**Approved cashier NFC location:** [INSERT LOCATION WITHIN LICENSED PREMISES]

The club is the offeror and will train cashier staff, display material restrictions before selection, protect the official MyDancr NFC tag, and honor each valid server-confirmed redemption subject only to the disclosed lawful conditions.

## 3. Negotiated referral fee

**MyDancr referral fee for each eligible verified redemption:** [INSERT FIXED FEE OR FORMULA AND CURRENCY]

The referral fee is earned only when the MyDancr server confirms the correct active deal through the official cashier NFC tag and the event is not expired, duplicate, fraudulent, reversed, prohibited, or otherwise ineligible. A dancer commission, when applicable, is funded and calculated separately under the Dancer Commission Terms and does not reduce the club's referral-fee obligation.

## 4. Weekly statement and ACH

**Statement week and cutoff:** [INSERT WEEK, CUTOFF TIME, AND ARIZONA TIMEZONE]

**Statement delivery:** [INSERT ACCOUNT AND EMAIL METHOD]

**ACH method:** [INSERT AUTHORIZED DEBIT OR CLUB-INITIATED CREDIT METHOD]

**Weekly debit or payment day:** [INSERT DAY AND HOLIDAY RULE]

**Banking authorization reference:** [INSERT SEPARATE ACH AUTHORIZATION OR PROCESSOR RECORD]

The club will pay the undisputed weekly statement by ACH once each week. The club authorizes MyDancr and its payment processor to use the approved method solely for amounts due under accepted Deal Orders, returns, corrections, and agreed fees. The club must keep valid banking instructions and promptly update changes. Returned or rejected ACH entries may be retried as permitted by the authorization and law. The club is responsible for bank-return charges and lawful collection costs stated in the final accepted order.

## 5. Review, disputes, and records

The club must dispute a statement item in writing within [INSERT NUMBER] days after delivery and identify the order, deal, redemption reference, amount, and basis. Undisputed amounts remain due. MyDancr will preserve the server redemption, NFC confirmation, attribution source, deal terms, fee snapshot, statement item, ACH status, corrections, and audit history for the applicable contract, accounting, tax, fraud, and limitations periods. A good-faith dispute does not authorize alteration of an NFC tag, attribution, or customer record.

## 6. Term, suspension, and acceptance

This Deal Order begins on the later of its electronic acceptance or stated offer start and continues until the stated end, replacement, or termination under the Club Agreement. MyDancr may pause publication, redemption, or promotion for nonpayment, expired terms, tag compromise, deceptive fulfillment, legal risk, or breach. A club must remain NFC-enabled and maintain at least one active approved Club Deal to receive promoted public listing and roster placement under the current business model.

**Club electronic acceptance:** The authorized representative accepts through the recorded MyDancr clickwrap or approved electronic signature process.

**MyDancr acceptance:** Acceptance is recorded by the authorized Company account or signature process.
""".strip()

    front = """
# Documents in this bundle

This bundle contains only the operative contract and policy drafts prepared for counsel to redline. It intentionally omits the product map, implementation checklist, test scripts, internal commentary, and source list from the broader legal review packet.

- Document 1 - Terms of Service
- Document 2 - Privacy Policy
- Document 3 - Dancer Agreement
- Schedule A - Content and Media Consent/License
- Schedule B - Acceptable Use and Prohibited Conduct Rules
- Schedule C - Account Eligibility and Venue Verification
- Schedule D - Dancer Commission and Too Much Media/NATS Settlement Terms
- Document 4 - Club Agreement
- Schedule A - Club Deal and NFC Redemption Terms
- Schedule B - Club Deal Order and Weekly ACH Schedule
- Document 5 - DMCA and Takedown Policy

Each document is an attorney review draft, not approved for publication or acceptance. The final versions must identify the contracting legal entity, notice contacts, effective dates, version numbers, and any deal-specific or provider-specific commercial fields.
""".strip()

    sections = [
        front,
        terms,
        privacy,
        dancer,
        media,
        acceptable,
        eligibility,
        commission,
        club,
        deal,
        normalize_ascii(deal_order),
        dmca,
    ]
    return "\n\n---PAGE---\n\n".join(sections) + "\n"


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 26, INK, 0, 8),
        ("Subtitle", 13, MUTED, 0, 8),
        ("Heading 1", 16, BLUE, 14, 8),
        ("Heading 2", 13, BLUE, 11, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style_ppr = style._element.get_or_add_pPr()
        style_border = style_ppr.find(qn("w:pBdr"))
        if style_border is not None:
            style_ppr.remove(style_border)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("MYDANCR | ATTORNEY REDLINE DRAFT")
    set_run_font(hr, size=8, color=MUTED, bold=True)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(5.5)
    table.columns[1].width = Inches(1)
    left = table.cell(0, 0).paragraphs[0]
    lr = left.add_run("Confidential working draft | Not approved for publication")
    set_run_font(lr, size=8, color=MUTED)
    set_page_number(table.cell(0, 1).paragraphs[0])
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is not None:
            tc_w.set(qn("w:type"), "dxa")


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(90)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("MYDANCR")
    set_run_font(kr, size=12, color=BLUE, bold=True)
    kicker.paragraph_format.space_after = Pt(16)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Attorney Redline Documents")
    set_run_font(tr, size=26, color=INK, bold=True)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Operative contract and policy drafts only")
    set_run_font(sr, size=13, color=MUTED, bold=False)
    subtitle.paragraph_format.space_after = Pt(42)

    for label, value in (
        ("Prepared for", "Specialist adult-entertainment counsel"),
        ("Pilot market", "Arizona"),
        ("Review date", "August 21, 2026"),
        ("Status", "Attorney review draft - not approved for publication or acceptance"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, size=10.5, color=INK, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.5, color=INK)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(32)
    note.paragraph_format.left_indent = Inches(0.55)
    note.paragraph_format.right_indent = Inches(0.55)
    nr = note.add_run("Prepared as a drafting aid for licensed counsel. This is not legal advice or a final legal opinion.")
    set_run_font(nr, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()


def parse_inline(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color=DARK_BLUE)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def render_markdown(doc: Document, source: str) -> None:
    lines = source.splitlines()
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if stripped == "---PAGE---":
            doc.add_page_break()
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, re.sub(r"^\d+\.\s", "", stripped))
        else:
            paragraph_lines = [stripped]
            j = i + 1
            while j < len(lines):
                candidate = lines[j].strip()
                if not candidate or candidate.startswith(("#", "- ", "---PAGE---")) or re.match(r"^\d+\.\s", candidate):
                    break
                paragraph_lines.append(candidate)
                j += 1
            p = doc.add_paragraph()
            parse_inline(p, " ".join(paragraph_lines))
            i = j - 1
        i += 1


def main() -> None:
    if len(sys.argv) != 4:
        raise SystemExit("Usage: create_attorney_redline_documents.py SOURCE_PACKET.md OUTPUT_SOURCE.md OUTPUT.docx")
    source_packet = Path(sys.argv[1]).resolve()
    output_source = Path(sys.argv[2]).resolve()
    output_docx = Path(sys.argv[3]).resolve()
    output_source.parent.mkdir(parents=True, exist_ok=True)
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    source = build_source(source_packet.read_text(encoding="utf-8"))
    output_source.write_text(source, encoding="utf-8")

    doc = Document()
    set_document_defaults(doc)
    add_cover(doc)
    render_markdown(doc, source)
    doc.core_properties.title = "MyDancr Attorney Redline Documents"
    doc.core_properties.subject = "Operative contract and policy drafts for attorney redline"
    doc.core_properties.author = "MyDancr product team"
    doc.core_properties.keywords = "MyDancr, attorney redline, terms, privacy, dancer agreement, club agreement, NFC, NATS, Arizona"
    doc.core_properties.comments = "Attorney review draft; not approved for publication or acceptance"
    doc.save(output_docx)
    print(output_docx)


if __name__ == "__main__":
    main()
